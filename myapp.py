import streamlit as st
from openai import OpenAI
import os
from datetime import datetime
import json
import io
import sys

# Import document parsing libraries with error handling
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    import pandas as pd
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

# Page configuration
st.set_page_config(
    page_title="Document Chat Assistant",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'About': "📄 Document Chat Assistant v2.0\nChat with PDF, DOCX, XLS, TXT, and MD files"
    }
)

# Custom CSS for better UI
st.markdown("""
    <style>
    .main { padding: 2rem 1rem; }
    .stChatMessage { margin: 1rem 0; }
    [data-testid="stSidebar"] { background-color: #f0f2f6; }
    .document-info { background: #e3f2fd; padding: 1rem; border-radius: 0.5rem; margin: 1rem 0; border-left: 4px solid #2196F3; }
    .chat-container { max-width: 900px; margin: 0 auto; }
    .file-type-badge { display: inline-block; padding: 0.25rem 0.75rem; background: #4CAF50; color: white; border-radius: 0.25rem; font-size: 0.8rem; margin-left: 0.5rem; font-weight: bold; }
    .error-box { background: #ffebee; padding: 1rem; border-radius: 0.5rem; color: #c62828; }
    .success-box { background: #e8f5e9; padding: 1rem; border-radius: 0.5rem; color: #2e7d32; }
    .info-box { background: #e3f2fd; padding: 1rem; border-radius: 0.5rem; color: #1565c0; }
    </style>
""", unsafe_allow_html=True)

# Initialize session state
if "messages" not in st.session_state:
    st.session_state.messages = []
if "document_content" not in st.session_state:
    st.session_state.document_content = None
if "document_name" not in st.session_state:
    st.session_state.document_name = None
if "document_type" not in st.session_state:
    st.session_state.document_type = None
if "document_size" not in st.session_state:
    st.session_state.document_size = 0
if "saved_contexts" not in st.session_state:
    st.session_state.saved_contexts = {}
if "client" not in st.session_state:
    st.session_state.client = None
if "show_save_dialog" not in st.session_state:
    st.session_state.show_save_dialog = False

@st.cache_resource
def initialize_openai_client(api_key):
    """Initialize and cache OpenAI client"""
    try:
        return OpenAI(api_key=api_key)
    except Exception as e:
        st.error(f"Failed to initialize OpenAI client: {e}")
        return None

@st.cache_data
def process_document(file_content: str, file_name: str, file_type: str):
    """Cache processed document content"""
    return {
        "content": file_content,
        "name": file_name,
        "type": file_type,
        "size": len(file_content),
        "processed_at": datetime.now().isoformat()
    }

def extract_pdf_content(file_object) -> str:
    """Extract text content from PDF file"""
    try:
        if not PYPDF2_AVAILABLE:
            raise Exception("PyPDF2 not installed. Install with: pip install PyPDF2")
        
        pdf_reader = PyPDF2.PdfReader(file_object)
        text = ""
        for page_num in range(len(pdf_reader.pages)):
            page = pdf_reader.pages[page_num]
            text += f"\n--- Page {page_num + 1} ---\n"
            extracted = page.extract_text()
            if extracted:
                text += extracted
        
        if not text.strip():
            raise Exception("No text found in PDF. The document may be image-based.")
        
        return text
    except Exception as e:
        raise Exception(f"Error extracting PDF: {str(e)}")

def extract_docx_content(file_object) -> str:
    """Extract text content from DOCX file"""
    try:
        if not DOCX_AVAILABLE:
            raise Exception("python-docx not installed. Install with: pip install python-docx")
        
        doc = Document(file_object)
        text = ""
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        
        # Extract tables if any
        if doc.tables:
            text += "\n[TABLES]\n"
            for table_idx, table in enumerate(doc.tables):
                text += f"\nTable {table_idx + 1}:\n"
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    text += row_text + "\n"
        
        if not text.strip():
            raise Exception("No content found in DOCX document.")
        
        return text
    except Exception as e:
        raise Exception(f"Error extracting DOCX: {str(e)}")

def extract_excel_content(file_object) -> str:
    """Extract content from Excel file (XLS/XLSX)"""
    try:
        if not EXCEL_AVAILABLE:
            raise Exception("pandas/openpyxl not installed. Install with: pip install pandas openpyxl")
        
        excel_file = pd.ExcelFile(file_object)
        text = ""
        
        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_object, sheet_name=sheet_name)
            text += f"\n--- Sheet: {sheet_name} ---\n"
            text += df.to_string(index=True)
            text += "\n"
        
        if not text.strip():
            raise Exception("No data found in Excel file.")
        
        return text
    except Exception as e:
        raise Exception(f"Error extracting Excel: {str(e)}")

def extract_text_content(file_object) -> str:
    """Extract text from TXT or MD files"""
    try:
        content = file_object.read().decode('utf-8')
        if not content.strip():
            raise Exception("File is empty.")
        return content
    except Exception as e:
        raise Exception(f"Error reading text file: {str(e)}")

def save_context(context_name: str):
    """Save current conversation context"""
    if not context_name.strip():
        st.error("Context name cannot be empty.")
        return
    
    if st.session_state.messages:
        st.session_state.saved_contexts[context_name] = {
            "messages": st.session_state.messages.copy(),
            "document": st.session_state.document_name,
            "document_type": st.session_state.document_type,
            "saved_at": datetime.now().isoformat()
        }
        st.success(f"✅ Context '{context_name}' saved successfully!")
    else:
        st.warning("No conversation to save.")

def load_context(context_name: str):
    """Load saved conversation context"""
    if context_name in st.session_state.saved_contexts:
        context = st.session_state.saved_contexts[context_name]
        st.session_state.messages = context["messages"].copy()
        st.success(f"✅ Context '{context_name}' loaded successfully!")
        st.rerun()
    else:
        st.error("Context not found.")

def clear_chat():
    """Clear current conversation"""
    st.session_state.messages = []
    st.rerun()

def delete_context(context_name: str):
    """Delete a saved context"""
    if context_name in st.session_state.saved_contexts:
        del st.session_state.saved_contexts[context_name]
        st.rerun()

# Sidebar configuration
with st.sidebar:
    st.title("⚙️ Configuration")
    
    # API Key input
    api_key = st.text_input(
        "OpenAI API Key",
        type="password",
        help="Enter your OpenAI API key. Get one at https://platform.openai.com/api-keys"
    )
    
    if api_key:
        st.session_state.client = initialize_openai_client(api_key)
        st.success("✅ API Key configured")
    else:
        st.warning("⚠️ No API Key provided")
    
    st.divider()
    
    # Document upload section
    st.subheader("📁 Document Upload")
    st.caption("Supported: TXT, MD, PDF, DOCX, XLS, XLSX")
    
    # Show available formats
    available_formats = []
    if PYPDF2_AVAILABLE:
        available_formats.append("✓ PDF")
    if DOCX_AVAILABLE:
        available_formats.append("✓ DOCX")
    if EXCEL_AVAILABLE:
        available_formats.append("✓ Excel")
    available_formats.extend(["✓ TXT", "✓ MD"])
    
    if available_formats:
        st.caption("Available: " + ", ".join(available_formats))
    
    uploaded_file = st.file_uploader(
        "Upload a document",
        type=["txt", "md", "pdf", "docx", "doc", "xls", "xlsx"],
        accept_multiple_files=False
    )
    
    if uploaded_file:
        try:
            file_extension = uploaded_file.name.split('.')[-1].lower()
            file_content = None
            
            # Reset file pointer to beginning
            uploaded_file.seek(0)
            
            with st.spinner(f"📂 Processing {file_extension.upper()} file..."):
                if file_extension == "pdf":
                    file_content = extract_pdf_content(uploaded_file)
                    doc_type = "PDF"
                elif file_extension in ["docx", "doc"]:
                    file_content = extract_docx_content(uploaded_file)
                    doc_type = "DOCX"
                elif file_extension in ["xls", "xlsx"]:
                    file_content = extract_excel_content(uploaded_file)
                    doc_type = "Excel"
                elif file_extension in ["txt", "md"]:
                    file_content = extract_text_content(uploaded_file)
                    doc_type = "Text"
                
                if file_content:
                    processed = process_document(file_content, uploaded_file.name, doc_type)
                    st.session_state.document_content = processed["content"]
                    st.session_state.document_name = processed["name"]
                    st.session_state.document_type = processed["type"]
                    st.session_state.document_size = processed["size"]
                    st.success(f"✅ Document loaded: {uploaded_file.name}")
        
        except Exception as e:
            st.error(f"❌ Error loading document: {e}")
    
    # Display current document info
    if st.session_state.document_name:
        doc_type_badge = f'<span class="file-type-badge">{st.session_state.document_type}</span>'
        size_kb = st.session_state.document_size / 1024
        st.markdown(f"""
        <div class="document-info">
        <strong>📄 Current Document:</strong><br>
        {st.session_state.document_name} {doc_type_badge}<br>
        <small>Size: {size_kb:.2f} KB</small>
        </div>
        """, unsafe_allow_html=True)
        
        # Clear document button
        if st.button("🔄 Clear Document", use_container_width=True):
            st.session_state.document_content = None
            st.session_state.document_name = None
            st.session_state.document_type = None
            st.session_state.document_size = 0
            st.rerun()
    
    st.divider()
    
    # Context management
    st.subheader("💾 Context Management")
    
    col1, col2 = st.columns(2)
    with col1:
        if st.button("💾 Save Context", use_container_width=True):
            st.session_state.show_save_dialog = True
    
    with col2:
        if st.button("🗑️ Clear Chat", use_container_width=True):
            clear_chat()
    
    # Save dialog
    if st.session_state.show_save_dialog:
        context_name = st.text_input("Context name:", key="context_name_input", placeholder="e.g., Project Discussion")
        col1, col2 = st.columns(2)
        with col1:
            if st.button("✅ Save", key="save_btn", use_container_width=True):
                if context_name:
                    save_context(context_name)
                    st.session_state.show_save_dialog = False
                    st.rerun()
                else:
                    st.error("Please enter a context name")
        with col2:
            if st.button("❌ Cancel", key="cancel_btn", use_container_width=True):
                st.session_state.show_save_dialog = False
                st.rerun()
    
    # Display saved contexts
    if st.session_state.saved_contexts:
        st.subheader("📋 Saved Contexts")
        for context_name in st.session_state.saved_contexts.keys():
            col1, col2 = st.columns([3, 1])
            with col1:
                if st.button(f"📂 {context_name}", use_container_width=True, key=f"load_{context_name}"):
                    load_context(context_name)
            with col2:
                if st.button("🗑️", key=f"delete_{context_name}", help="Delete context"):
                    delete_context(context_name)
    
    st.divider()
    st.caption("📄 Document Chat Assistant v2.0")
    st.caption("Powered by OpenAI GPT-4")

# Main chat interface
st.title("📄 Document Chat Assistant")

# Check API key
if not st.session_state.client:
    st.warning("⚠️ Please enter your OpenAI API key in the sidebar to proceed.")
    st.info("Get your API key at: https://platform.openai.com/api-keys")
else:
    # Check document
    if not st.session_state.document_content:
        col1, col2, col3 = st.columns(3)
        with col1:
            st.info("📤 Upload a document in the sidebar to start chatting.")
        with col2:
            st.info("💡 Supported formats: PDF, DOCX, XLS, TXT, MD")
        with col3:
            st.info("🔗 Ask questions about your document")
    else:
        # Display chat messages
        st.subheader(f"Chat about: {st.session_state.document_name}")
        
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
        
        # Chat input
        if prompt := st.chat_input("Ask something about your document..."):
            if not st.session_state.document_content:
                st.error("Please upload a document first.")
            else:
                # Add user message
                st.session_state.messages.append({"role": "user", "content": prompt})
                
                with st.chat_message("user"):
                    st.markdown(prompt)
                
                # Generate response with streaming
                with st.chat_message("assistant"):
                    try:
                        system_prompt = f"""You are a helpful and professional assistant. Answer questions based ONLY on the following document content.

DOCUMENT: {st.session_state.document_name}
TYPE: {st.session_state.document_type}

DOCUMENT CONTENT:
---
{st.session_state.document_content}
---

Guidelines:
- Be concise and helpful
- If the answer is not in the document, clearly say so
- Provide relevant context from the document when answering
- If asked to do something outside the document scope, politely decline"""
                        
                        with st.spinner("🤖 Thinking..."):
                            stream = st.session_state.client.chat.completions.create(
                                model="gpt-4o",
                                messages=[
                                    {"role": "system", "content": system_prompt},
                                    *st.session_state.messages
                                ],
                                stream=True,
                                temperature=0.7,
                                max_tokens=1000,
                                top_p=0.95
                            )
                            
                            response_content = st.write_stream(stream)
                        
                        # Add assistant message
                        st.session_state.messages.append({"role": "assistant", "content": response_content})
                    
                    except Exception as e:
                        st.error(f"❌ Error generating response: {e}")

